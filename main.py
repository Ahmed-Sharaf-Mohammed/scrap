import re
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

# Set up the Selenium WebDriver
driver = webdriver.Chrome()

# Open the webpage
driver.get("http://e-books.helwan.edu.eg/storage/29946/index.html#/reader/chapter/32")

try:
    # Wait for the div element with the class 'WordSection2' to load
    specific_div = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.CLASS_NAME, "WordSection2"))
    )
    # Get the HTML content
    html_content = specific_div.get_attribute('innerHTML')

    # Parse the HTML with BeautifulSoup
    soup = BeautifulSoup(html_content, "html.parser")

    # Create a new Word Document
    doc = Document()

    # Function to clean unwanted text
    def clean_text(text):
        # Remove [if !supportLists], [endif], and other similar patterns
        return re.sub(r'\[.*?\]', '', text).strip()

    # A set to track processed texts to avoid duplicates
    processed_texts = set()

    # Extract and format content
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
        # If it's a heading, apply a heading style
        if element.name.startswith('h'):
            heading_text = clean_text(element.get_text(separator=" ", strip=True))
            if heading_text and heading_text not in processed_texts:  # Only add if text is not empty
                doc.add_paragraph(heading_text, style='Heading ' + element.name[1])
                processed_texts.add(heading_text)  # Mark as processed

        # If it's a paragraph, create a paragraph and add text
        elif element.name == 'p':
            paragraph = doc.add_paragraph()
            for child in element.descendants:
                if child.string and child.string.strip():  # Check if it's a valid text node
                    text = clean_text(child.string.strip())
                    if not text:  # Skip if the cleaned text is empty or already processed
                        continue

                    # Check if the text has already been processed
                    if text in processed_texts:
                        continue

                    # Create a new run for the text
                    run = paragraph.add_run(text + " ")

                    # Set of styles to be applied on the run (prevent applying the same style multiple times)
                    applied_styles = set()

                    # Apply styles based on parent tags (only once per style)
                    for parent in child.parents:
                        if parent.name in ['b', 'strong'] and 'bold' not in applied_styles:
                            run.bold = True
                            applied_styles.add('bold')  # Mark 'bold' style as applied
                        if parent.name in ['i', 'em'] and 'italic' not in applied_styles:
                            run.italic = True
                            applied_styles.add('italic')  # Mark 'italic' style as applied
                        if parent.name == 'mark' and 'highlight' not in applied_styles:
                            run.font.highlight_color = 7  # Highlight yellow
                            applied_styles.add('highlight')  # Mark 'highlight' style as applied
                        if parent.name == 'small' and 'small' not in applied_styles:
                            run.font.size = Pt(8)  # Smaller font
                            applied_styles.add('small')  # Mark 'small' style as applied
                        if parent.name == 'del' and 'strike' not in applied_styles:
                            run.font.strike = True  # Strikethrough
                            applied_styles.add('strike')  # Mark 'strike' style as applied
                        if parent.name == 'ins' and 'underline' not in applied_styles:
                            run.underline = True  # Underline
                            applied_styles.add('underline')  # Mark 'underline' style as applied
                        if parent.name == 'u' and 'underline' not in applied_styles:
                            run.underline = True  # Underline
                            applied_styles.add('underline')  # Mark 'underline' style as applied
                        if parent.name == 'sup' and 'superscript' not in applied_styles:
                            run.font.superscript = True  # Superscript
                            applied_styles.add('superscript')  # Mark 'superscript' style as applied

                    # Debug: print the text and styles being applied to it
                    print(f"Text: '{text}'")
                    print(f"Applied Styles: {applied_styles}")

                    # Mark this text as processed
                    processed_texts.add(text)

    # Save the Word document
    doc.save("cleaned_extracted_content_with_headings.docx")

except Exception as e:
    print(f"An error occurred: {e}")

finally:
    # Close the browser
    driver.quit()
